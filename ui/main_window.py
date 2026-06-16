import re
import os
import json
from PyQt5.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QLineEdit, QTextEdit, QListWidget, QListWidgetItem, QSplitter,
    QFrame, QComboBox, QMessageBox, QFileDialog, QDialog, QApplication,
    QAbstractItemView
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QMimeData, QUrl
from PyQt5.QtGui import QFont, QColor, QPixmap, QDrag, QPainter

from core.scanner import scan_folder
from core.search import search_documents
from core.date_extractor import check_expiry_from_excel
from core.summarizer import summarize_document, generate_document
from docx import Document as DocxDocument

AI_HISTORY_FILE = "cache/ai_history.json"


class WorkerThread(QThread):
    finished = pyqtSignal(str)

    def __init__(self, func, *args):
        super().__init__()
        self.func = func
        self.args = args

    def run(self):
        result = self.func(*self.args)
        self.finished.emit(str(result))


THEME = {
    "bg": "#f4f7fb",
    "sidebar": "#ffffff",
    "surface": "#ffffff",
    "surface_2": "#f9fbff",
    "input": "#ffffff",
    "border": "#d9e1ec",
    "accent": "#4f7cff",
    "accent_soft": "#e9efff",
    "text": "#1f2937",
    "muted": "#6b7280",
    "success": "#1f9d73",
    "danger": "#e35d6a",
}


class WorkerPanel(QFrame):
    def __init__(self):
        super().__init__()
        self.setObjectName("workerPanel")


class InfoCard(QFrame):
    def __init__(self, title, desc):
        super().__init__()
        self.setObjectName("infoCard")

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(8)

        self.title_label = QLabel(title)
        self.title_label.setFont(QFont("Segoe UI Semibold", 11))

        self.desc_label = QLabel(desc)
        self.desc_label.setWordWrap(True)
        self.desc_label.setFont(QFont("Segoe UI", 10))

        layout.addWidget(self.title_label)
        layout.addWidget(self.desc_label)
        layout.addStretch()


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("DocuMind AI")
        self.setMinimumSize(1240, 760)

        self.theme = THEME
        self.documents = []
        self.search_results = []
        self.sum_filtered_docs = []
        self.last_generated = ""
        self.worker = None
        self.status_level = "muted"

        self.init_ui()
        self.apply_theme()
        self._set_logo_mark_icon()
        self.show_panel("search")

    # --------------------------------------------------
    # UI
    # --------------------------------------------------
    def init_ui(self):
        root = QWidget()
        self.setCentralWidget(root)

        root_layout = QHBoxLayout(root)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(0)

        # SIDEBAR
        self.sidebar = QFrame()
        self.sidebar.setFixedWidth(250)

        sidebar_layout = QVBoxLayout(self.sidebar)
        sidebar_layout.setContentsMargins(18, 18, 18, 18)
        sidebar_layout.setSpacing(18)

        brand_wrap = QWidget()
        brand_layout = QVBoxLayout(brand_wrap)
        brand_layout.setContentsMargins(0, 0, 0, 0)
        brand_layout.setSpacing(4)

        brand_top = QHBoxLayout()
        brand_top.setContentsMargins(0, 0, 0, 0)
        brand_top.setSpacing(10)

        self.logo_mark = QLabel()
        self.logo_mark.setAlignment(Qt.AlignCenter)
        self.logo_mark.setFixedSize(42, 42)
        self.logo_mark.setFont(QFont("Segoe UI Semibold", 11))

        logo_text_wrap = QWidget()
        logo_text_layout = QVBoxLayout(logo_text_wrap)
        logo_text_layout.setContentsMargins(0, 0, 0, 0)
        logo_text_layout.setSpacing(1)

        self.logo = QLabel("DocuMind AI")
        self.logo.setFont(QFont("Segoe UI Semibold", 15))

        self.logo_subtitle = QLabel("Smart document workspace")
        self.logo_subtitle.setFont(QFont("Segoe UI", 9))

        logo_text_layout.addWidget(self.logo)
        logo_text_layout.addWidget(self.logo_subtitle)

        brand_top.addWidget(self.logo_mark)
        brand_top.addWidget(logo_text_wrap)

        brand_layout.addLayout(brand_top)
        sidebar_layout.addWidget(brand_wrap)

        self.sidebar_divider_1 = QFrame()
        self.sidebar_divider_1.setFixedHeight(1)
        sidebar_layout.addWidget(self.sidebar_divider_1)

        self.tabs = {}
        nav_items = [
            ("search", "Search"),
            ("deadline", "Deadlines"),
            ("summarize", "Summarize"),
            ("generate", "Generate"),
        ]

        nav_wrap = QWidget()
        nav_layout = QVBoxLayout(nav_wrap)
        nav_layout.setContentsMargins(0, 0, 0, 0)
        nav_layout.setSpacing(8)

        for key, text in nav_items:
            btn = QPushButton(text)
            btn.setCursor(Qt.PointingHandCursor)
            btn.setFixedHeight(42)
            btn.setFont(QFont("Segoe UI Semibold", 10))
            btn.clicked.connect(lambda _, k=key: self.show_panel(k))
            self.tabs[key] = btn
            nav_layout.addWidget(btn)

        sidebar_layout.addWidget(nav_wrap)
        sidebar_layout.addStretch()

        self.folder_btn = QPushButton("Select Folder")
        self.folder_btn.setCursor(Qt.PointingHandCursor)
        self.folder_btn.setFixedHeight(42)
        self.folder_btn.clicked.connect(self.select_folder)
        sidebar_layout.addWidget(self.folder_btn)

        self.folder_name = QLabel("No folder selected yet")
        self.folder_name.setWordWrap(True)
        self.folder_name.setFont(QFont("Segoe UI", 9))
        sidebar_layout.addWidget(self.folder_name)

        root_layout.addWidget(self.sidebar)

        # MAIN AREA
        self.main_area = QWidget()
        main_layout = QVBoxLayout(self.main_area)
        main_layout.setContentsMargins(22, 18, 22, 18)
        main_layout.setSpacing(18)

        # HEADER
        self.header = QWidget()
        header_layout = QHBoxLayout(self.header)
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(12)

        title_wrap = QWidget()
        title_layout = QVBoxLayout(title_wrap)
        title_layout.setContentsMargins(0, 0, 0, 0)
        title_layout.setSpacing(2)

        self.page_title = QLabel("Search")
        self.page_title.setFont(QFont("Segoe UI Semibold", 22))

        self.page_subtitle = QLabel("Quickly find information inside your scanned documents")
        self.page_subtitle.setFont(QFont("Segoe UI", 10))

        title_layout.addWidget(self.page_title)
        title_layout.addWidget(self.page_subtitle)

        header_layout.addWidget(title_wrap)
        header_layout.addStretch()

        self.sum_history_btn = QPushButton("History")
        self.sum_history_btn.setCursor(Qt.PointingHandCursor)
        self.sum_history_btn.setFixedHeight(34)
        self.sum_history_btn.clicked.connect(self.open_summarize_history_dialog)
        self.sum_history_btn.setVisible(False)
        header_layout.addWidget(self.sum_history_btn)

        self.header_status = QLabel("Ready to work")
        self.header_status.setFont(QFont("Segoe UI", 10))
        header_layout.addWidget(self.header_status)

        main_layout.addWidget(self.header)

        self.search_panel = self._build_search_panel()
        self.deadline_panel = self._build_deadline_panel()
        self.summarize_panel = self._build_summarize_panel()
        self.generate_panel = self._build_generate_panel()

        main_layout.addWidget(self.search_panel)
        main_layout.addWidget(self.deadline_panel)
        main_layout.addWidget(self.summarize_panel)
        main_layout.addWidget(self.generate_panel)

        self.statusbar_frame = QFrame()
        self.statusbar_frame.setFixedHeight(34)
        sb_layout = QHBoxLayout(self.statusbar_frame)
        sb_layout.setContentsMargins(0, 0, 0, 0)

        self.status = QLabel("Ready - select a folder to begin")
        self.status.setFont(QFont("Segoe UI", 9))

        self.footer_folder = QLabel("No folder selected yet")
        self.footer_folder.setFont(QFont("Segoe UI", 9))

        sb_layout.addWidget(self.status)
        sb_layout.addStretch()
        sb_layout.addWidget(self.footer_folder)

        main_layout.addWidget(self.statusbar_frame)

        root_layout.addWidget(self.main_area)

    def _build_search_panel(self):
        panel = WorkerPanel()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        search_row = QHBoxLayout()
        search_row.setSpacing(10)

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Try: contract date, invoice number, employee name...")
        self.search_input.setFixedHeight(44)
        self.search_input.returnPressed.connect(self.search)

        self.search_btn = QPushButton("Search")
        self.search_btn.setCursor(Qt.PointingHandCursor)
        self.search_btn.setFixedHeight(44)
        self.search_btn.setFixedWidth(130)
        self.search_btn.clicked.connect(self.search)

        search_row.addWidget(self.search_input)
        search_row.addWidget(self.search_btn)
        layout.addLayout(search_row)

        # FILL EMPTY SPACE
        cards_row = QHBoxLayout()
        cards_row.setSpacing(12)

        self.card_fast = InfoCard(
            "Fast Search",
            "Find matching text across all scanned documents in seconds."
        )
        self.card_private = InfoCard(
            "Private Workspace",
            "Your files are processed locally inside your own application."
        )
        self.card_formats = InfoCard(
            "Multiple Formats",
            "Search and review PDF, DOCX and XLSX documents in one place."
        )

        cards_row.addWidget(self.card_fast)
        cards_row.addWidget(self.card_private)
        cards_row.addWidget(self.card_formats)

        layout.addLayout(cards_row)

        split = QSplitter(Qt.Horizontal)
        split.setHandleWidth(1)
        split.setChildrenCollapsible(False)

        self.results_panel = QFrame()
        rp_layout = QVBoxLayout(self.results_panel)
        rp_layout.setContentsMargins(16, 14, 16, 16)
        rp_layout.setSpacing(10)

        results_header = QHBoxLayout()
        results_header.setContentsMargins(0, 0, 0, 0)
        results_header.setSpacing(8)

        self.results_label = QLabel("Matching Documents")
        self.results_label.setFont(QFont("Segoe UI Semibold", 11))

        self.results_count = QLabel("0")
        self.results_count.setAlignment(Qt.AlignCenter)
        self.results_count.setFixedSize(28, 22)
        self.results_count.setFont(QFont("Segoe UI Semibold", 9))

        results_header.addWidget(self.results_label)
        results_header.addWidget(self.results_count)
        results_header.addStretch()
        rp_layout.addLayout(results_header)

        self.results_list = FileCopyDragListWidget()
        self.results_list.itemClicked.connect(self.show_snippet)
        rp_layout.addWidget(self.results_list)

        self.results_empty = QLabel("No results yet. Run a search to see matching files.")
        self.results_empty.setAlignment(Qt.AlignCenter)
        self.results_empty.setFont(QFont("Segoe UI", 11))
        rp_layout.addWidget(self.results_empty)

        self.preview_panel = QFrame()
        pp_layout = QVBoxLayout(self.preview_panel)
        pp_layout.setContentsMargins(16, 14, 16, 16)
        pp_layout.setSpacing(10)

        preview_header = QHBoxLayout()
        preview_header.setContentsMargins(0, 0, 0, 0)
        preview_header.setSpacing(8)

        self.preview_label = QLabel("Preview")
        self.preview_label.setFont(QFont("Segoe UI Semibold", 11))

        preview_header.addWidget(self.preview_label)
        preview_header.addStretch()
        pp_layout.addLayout(preview_header)

        self.snippet_box = QTextEdit()
        self.snippet_box.setReadOnly(True)
        pp_layout.addWidget(self.snippet_box)

        self.preview_empty = QLabel("Select a result to preview the matching text.")
        self.preview_empty.setAlignment(Qt.AlignCenter)
        self.preview_empty.setFont(QFont("Segoe UI", 11))
        pp_layout.addWidget(self.preview_empty)

        split.addWidget(self.results_panel)
        split.addWidget(self.preview_panel)
        split.setSizes([360, 760])

        layout.addWidget(split)
        return panel

    def _build_deadline_panel(self):
        panel = WorkerPanel()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        self.deadline_list = FileCopyDragListWidget()
        layout.addWidget(self.deadline_list)

        return panel

    def _build_summarize_panel(self):
        panel = WorkerPanel()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        top = QLabel("Choose a Document")
        top.setFont(QFont("Segoe UI Semibold", 11))
        layout.addWidget(top)

        self.sum_search_input = QLineEdit()
        self.sum_search_input.setPlaceholderText("Search document in list...")
        self.sum_search_input.setFixedHeight(40)
        self.sum_search_input.textChanged.connect(self.filter_summarize_documents)
        layout.addWidget(self.sum_search_input)

        self.sum_doc_list = FileCopyDragListWidget()
        self.sum_doc_list.setFixedHeight(210)
        layout.addWidget(self.sum_doc_list)

        self.sum_btn = QPushButton("Summarize")
        self.sum_btn.setCursor(Qt.PointingHandCursor)
        self.sum_btn.setFixedHeight(44)
        self.sum_btn.clicked.connect(self.summarize_selected)
        layout.addWidget(self.sum_btn)

        self.sum_output = QTextEdit()
        self.sum_output.setReadOnly(True)
        layout.addWidget(self.sum_output)

        return panel

    def _build_generate_panel(self):
        panel = WorkerPanel()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        type_label = QLabel("Select Document Type")
        type_label.setFont(QFont("Segoe UI Semibold", 11))
        layout.addWidget(type_label)

        self.doc_type = QComboBox()
        self.doc_type.addItems([
            "Employment Contract",
            "Non-Disclosure Agreement (NDA)",
            "Invoice",
            "Custom",
        ])
        self.doc_type.setFixedHeight(44)
        layout.addWidget(self.doc_type)

        details_label = QLabel("Enter Details")
        details_label.setFont(QFont("Segoe UI Semibold", 11))
        layout.addWidget(details_label)

        self.gen_input = QLineEdit()
        self.gen_input.setPlaceholderText(
            "e.g. position: IT Support | company: ABC LLC | salary: 3000 GEL"
        )
        self.gen_input.setFixedHeight(44)
        layout.addWidget(self.gen_input)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(10)

        self.gen_btn = QPushButton("Generate Draft")
        self.gen_btn.setCursor(Qt.PointingHandCursor)
        self.gen_btn.setFixedHeight(44)
        self.gen_btn.clicked.connect(self.generate_doc)

        self.save_btn = QPushButton("Export to Word")
        self.save_btn.setCursor(Qt.PointingHandCursor)
        self.save_btn.setFixedHeight(44)
        self.save_btn.clicked.connect(self.save_as_word)

        btn_row.addWidget(self.gen_btn)
        btn_row.addWidget(self.save_btn)
        layout.addLayout(btn_row)

        self.gen_output = QTextEdit()
        self.gen_output.setReadOnly(True)
        layout.addWidget(self.gen_output)

        return panel

    # --------------------------------------------------
    # THEME
    # --------------------------------------------------
    def apply_theme(self):
        t = self.theme

        self.setStyleSheet(f"""
            QMainWindow {{
                background: {t['bg']};
            }}
            QWidget {{
                background: transparent;
                color: {t['text']};
            }}
            QToolTip {{
                background-color: {t['surface']};
                color: {t['text']};
                border: 1px solid {t['border']};
                padding: 6px;
            }}
        """)

        self.sidebar.setStyleSheet(f"""
            QFrame {{
                background: {t['sidebar']};
                border-right: 1px solid {t['border']};
            }}
        """)

        self.logo_mark.setStyleSheet(f"""
            background: {t['accent_soft']};
            color: {t['accent']};
            border: 1px solid #d7e2ff;
            border-radius: 14px;
        """)
        self.logo.setStyleSheet(f"color: {t['text']};")
        self.logo_subtitle.setStyleSheet(f"color: {t['muted']};")
        self.folder_name.setStyleSheet(f"color: {t['muted']};")
        self.sidebar_divider_1.setStyleSheet(f"background: {t['border']};")

        for key, btn in self.tabs.items():
            active = key == getattr(self, "_active_panel", "search")
            if active:
                btn.setStyleSheet(f"""
                    QPushButton {{
                        background: {t['accent_soft']};
                        color: {t['accent']};
                        border: 1px solid #cfdbff;
                        border-radius: 10px;
                        text-align: left;
                        padding-left: 14px;
                    }}
                """)
            else:
                btn.setStyleSheet(f"""
                    QPushButton {{
                        background: transparent;
                        color: {t['muted']};
                        border: none;
                        border-radius: 10px;
                        text-align: left;
                        padding-left: 14px;
                    }}
                    QPushButton:hover {{
                        background: {t['surface']};
                        color: {t['text']};
                    }}
                """)

        self.folder_btn.setStyleSheet(f"""
            QPushButton {{
                background: {t['accent']};
                color: white;
                border: none;
                border-radius: 12px;
                font-weight: 600;
            }}
            QPushButton:hover {{
                background: #3f6cf0;
            }}
        """)

        self.main_area.setStyleSheet(f"background: {t['bg']};")
        self.page_title.setStyleSheet(f"color: {t['text']};")
        self.page_subtitle.setStyleSheet(f"color: {t['muted']};")
        self.header_status.setStyleSheet(f"color: {t['muted']};")

        for panel in [self.search_panel, self.deadline_panel, self.summarize_panel, self.generate_panel]:
            panel.setStyleSheet(f"""
                QFrame#workerPanel {{
                    background: {t['surface']};
                    border: 1px solid {t['border']};
                    border-radius: 18px;
                }}
            """)

        inner_panel_style = """
            QFrame {
                background: transparent;
                border: none;
            }
        """
        self.results_panel.setStyleSheet(inner_panel_style)
        self.preview_panel.setStyleSheet(inner_panel_style)

        for card in [self.card_fast, self.card_private, self.card_formats]:
            card.setStyleSheet(f"""
                QFrame#infoCard {{
                    background: {t['surface_2']};
                    border: 1px solid {t['border']};
                    border-radius: 12px;
                }}
            """)
            card.title_label.setStyleSheet(f"color: {t['text']};")
            card.desc_label.setStyleSheet(f"color: {t['muted']};")

        self.search_input.setStyleSheet(f"""
            QLineEdit {{
                background: {t['input']};
                color: {t['text']};
                border: 1px solid {t['border']};
                border-radius: 12px;
                padding: 0 14px;
                font-size: 13px;
            }}
            QLineEdit:focus {{
                border: 1px solid {t['accent']};
            }}
        """)
        self.sum_search_input.setStyleSheet(self.search_input.styleSheet())
        self.gen_input.setStyleSheet(self.search_input.styleSheet())

        accent_btn = f"""
            QPushButton {{
                background: {t['accent']};
                color: white;
                border: none;
                border-radius: 12px;
                font-size: 13px;
                font-weight: 600;
                padding: 0 16px;
            }}
            QPushButton:hover {{
                background: #3f6cf0;
            }}
        """
        self.search_btn.setStyleSheet(accent_btn)
        self.sum_btn.setStyleSheet(accent_btn)
        self.gen_btn.setStyleSheet(accent_btn)
        self.sum_history_btn.setStyleSheet(f"""
            QPushButton {{
                background: {t['surface_2']};
                color: {t['text']};
                border: 1px solid {t['border']};
                border-radius: 10px;
                font-size: 12px;
                font-weight: 600;
                padding: 0 12px;
            }}
            QPushButton:hover {{
                background: {t['accent_soft']};
                color: {t['accent']};
            }}
        """)

        self.save_btn.setStyleSheet(f"""
            QPushButton {{
                background: {t['surface_2']};
                color: {t['text']};
                border: 1px solid {t['border']};
                border-radius: 12px;
                font-size: 13px;
                font-weight: 600;
                padding: 0 16px;
            }}
            QPushButton:hover {{
                background: {t['accent_soft']};
            }}
        """)

        list_style = f"""
            QListWidget {{
            background: {t['input']};
            color: {t['text']};
            border: 1px solid {t['border']};
            border-radius: 12px;
            padding: 8px;
            outline: none;
            font-size: 13px;
            font-family: Segoe UI;
        }}
        QListWidget::item {{
            padding: 10px 14px;
            margin: 3px 0;
            border-radius: 8px;
            border-left: 3px solid transparent;
        }}
        QListWidget::item:selected {{
            background: {t['accent_soft']};
            color: {t['accent']};
            border-left: 3px solid {t['accent']};
            font-weight: bold;
        }}
        QListWidget::item:hover:!selected {{
            background: #f0f4ff;
            border-left: 3px solid #c5d5ff;
        }}
        QScrollBar:horizontal {{
            height: 6px;
            background: {t['border']};
            border-radius: 3px;
            margin: 0px;
        }}
        QScrollBar::handle:horizontal {{
            background: {t['accent']};
            border-radius: 3px;
            min-width: 30px;
        }}
        QScrollBar::handle:horizontal:hover {{
            background: {t['accent_soft']};
        }}
        QScrollBar::add-line:horizontal,
        QScrollBar::sub-line:horizontal {{
            width: 0px;
        }}
        QScrollBar:vertical {{
            width: 6px;
            background: {t['border']};
            border-radius: 3px;
            margin: 0px;
        }}
        QScrollBar::handle:vertical {{
            background: {t['accent']};
            border-radius: 3px;
            min-height: 30px;
        }}
        QScrollBar::handle:vertical:hover {{
            background: {t['accent_soft']};
        }}
        QScrollBar::add-line:vertical,
        QScrollBar::sub-line:vertical {{
            height: 0px;
        }}
        """
        self.results_list.setStyleSheet(list_style)
        self.deadline_list.setStyleSheet(list_style)
        self.sum_doc_list.setStyleSheet(list_style)

        text_style = f"""
            QTextEdit {{
                background: {t['input']};
                color: {t['text']};
                border: 1px solid {t['border']};
                border-radius: 12px;
                padding: 12px;
                font-size: 13px;
            }}
        """
        self.snippet_box.setStyleSheet(text_style)
        self.sum_output.setStyleSheet(text_style)
        self.gen_output.setStyleSheet(text_style)

        self.doc_type.setStyleSheet(f"""
            QComboBox {{
                background: {t['input']};
                color: {t['text']};
                border: 1px solid {t['border']};
                border-radius: 12px;
                padding: 0 12px;
                font-size: 13px;
            }}
            QComboBox::drop-down {{
                border: none;
                width: 28px;
            }}
            QComboBox QAbstractItemView {{
                background: {t['surface']};
                color: {t['text']};
                border: 1px solid {t['border']};
                selection-background-color: {t['accent_soft']};
                selection-color: {t['accent']};
            }}
        """)

        self.results_label.setStyleSheet(f"""
            color: {t['text']};
            background: transparent;
            border: none;
        """)

        self.preview_label.setStyleSheet(f"""
            color: {t['muted']};
            background: transparent;
            border: none;
        """)

        self.results_count.setStyleSheet(f"""
            background: transparent;
            color: {t['muted']};
            border: 1px solid {t['border']};
            border-radius: 7px;
            padding: 0 4px;
        """)

        self.results_empty.setStyleSheet(f"color: {t['muted']};")
        self.preview_empty.setStyleSheet(f"color: {t['muted']};")

        self.statusbar_frame.setStyleSheet(f"""
            background: transparent;
            border: none;
            border-top: 1px solid {t['border']};
            border-radius: 0px;
        """)

        self.status.setStyleSheet(f"""
            color: {t['muted']};
            background: transparent;
            border: none;
        """)
        self.apply_status_style()

        self.footer_folder.setStyleSheet(f"""
            color: {t['muted']};
            background: transparent;
            border: none;
        """)

        self.update_empty_states()

    # --------------------------------------------------
    # PANEL STATE
    # --------------------------------------------------
    def show_panel(self, name):
        self._active_panel = name

        self.search_panel.setVisible(name == "search")
        self.deadline_panel.setVisible(name == "deadline")
        self.summarize_panel.setVisible(name == "summarize")
        self.generate_panel.setVisible(name == "generate")

        titles = {
            "search": ("Search", "Search inside your scanned documents"),
            "deadline": ("Deadlines", "Upcoming expiry dates from spreadsheet files"),
            "summarize": ("Summarize", "Generate document summaries with AI"),
            "generate": ("Generate", "Create document templates with AI"),
        }
        self.page_title.setText(titles[name][0])
        self.page_subtitle.setText(titles[name][1])
        self.sum_history_btn.setVisible(name == "summarize")

        self.apply_theme()

    def _set_logo_mark_icon(self):
        """
        If no external logo image is provided, draw a small icon inside logo_mark.
        This keeps the UI looking consistent.
        """
        size = 42
        t = self.theme

        pix = QPixmap(size, size)
        pix.fill(Qt.transparent)

        painter = QPainter(pix)
        painter.setRenderHint(QPainter.Antialiasing, True)

        # Rounded background (matches the label's theme style).
        painter.setPen(QColor("#d7e2ff"))
        painter.setBrush(QColor(t["accent_soft"]))
        painter.drawRoundedRect(0, 0, size, size, 14, 14)

        # Center mark.
        painter.setPen(QColor(t["accent"]))
        painter.setFont(QFont("Segoe UI Semibold", 14, QFont.Bold))
        painter.drawText(pix.rect(), Qt.AlignCenter, "DM")
        painter.end()

        self.logo_mark.setPixmap(pix)

    # --------------------------------------------------
    # HELPERS
    # --------------------------------------------------
    def update_empty_states(self):
        has_results = len(self.search_results) > 0
        self.results_list.setVisible(has_results)
        self.results_empty.setVisible(not has_results)

        has_preview = bool(self.snippet_box.toPlainText().strip())
        self.snippet_box.setVisible(has_preview)
        self.preview_empty.setVisible(not has_preview)

    def set_results_count(self, count):
        self.results_count.setText(str(count))

    def refresh_summarize_documents(self):
        query = self.sum_search_input.text().strip().casefold()
        self.sum_doc_list.clear()

        self.sum_filtered_docs = []
        for doc in self.documents:
            filename = doc.get("filename", "")
            if query and query not in filename.casefold():
                continue
            self.sum_filtered_docs.append(doc)
            item = QListWidgetItem(filename)
            item.setData(Qt.UserRole, doc)
            self.sum_doc_list.addItem(item)

    def filter_summarize_documents(self):
        self.refresh_summarize_documents()

    def load_ai_history(self):
        if not os.path.exists(AI_HISTORY_FILE):
            return []
        try:
            with open(AI_HISTORY_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                return data
        except Exception:
            return []
        return []

    def _format_history_entry(self, entry):
        payload_preview = entry.get("payload_preview", "")
        response_preview = entry.get("response_preview", "")
        model = entry.get("model", "")
        cache_key = entry.get("cache_key", "")
        ts = entry.get("timestamp", "")
        task = entry.get("task", "")
        source = entry.get("source", "")

        return (
            f"Timestamp: {ts}\n"
            f"Task: {task}\n"
            f"Source: {source}\n"
            f"Model: {model}\n"
            f"Cache Key: {cache_key}\n\n"
            f"Input Preview:\n{payload_preview}\n\n"
            f"Output Preview:\n{response_preview}"
        )

    def open_summarize_history_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Summarizer History")
        dialog.resize(820, 560)

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        history_list = QListWidget()
        preview = QTextEdit()
        preview.setReadOnly(True)
        preview.setPlaceholderText("Select history item to preview details.")
        history_list.setStyleSheet(self.sum_doc_list.styleSheet())
        preview.setStyleSheet(self.sum_output.styleSheet())

        history_items = self.load_ai_history()
        for entry in reversed(history_items[-100:]):
            task = entry.get("task", "unknown")
            source = entry.get("source", "-")
            ts = entry.get("timestamp", "")
            label = f"{task}  |  {source}  |  {ts}"
            item = QListWidgetItem(label)
            item.setData(Qt.UserRole, entry)
            history_list.addItem(item)

        def on_item_clicked(item):
            entry = item.data(Qt.UserRole) or {}
            preview.setPlainText(self._format_history_entry(entry))

        history_list.itemClicked.connect(on_item_clicked)

        layout.addWidget(history_list)
        layout.addWidget(preview)
        dialog.exec_()

    def apply_status_style(self):
        color_map = {
            "muted": self.theme["muted"],
            "success": self.theme["success"],
            "error": self.theme["danger"],
        }
        color = color_map.get(self.status_level, self.theme["muted"])
        self.status.setStyleSheet(f"""
            color: {color};
            background: transparent;
            border: none;
            font-weight: 600;
        """)

    def set_status_message(self, text, level="muted"):
        self.status.setText(text)
        self.status_level = level
        self.apply_status_style()

    def show_warning(self, text):
        self.set_status_message(text, "error")
        box = QMessageBox(self)
        box.setIcon(QMessageBox.Warning)
        box.setWindowTitle("Action Required")
        box.setText(text)
        box.setStandardButtons(QMessageBox.Ok)
        ok_btn = box.button(QMessageBox.Ok)
        ok_btn.setText("OK")
        ok_btn.setStyleSheet("""
            QPushButton {
                min-width: 90px;
                min-height: 34px;
                border-radius: 10px;
                border: 1px solid #d1d5db;
                background-color: #ffffff;
                color: #111111;
                font-weight: 600;
                padding: 0 12px;
            }
            QPushButton:hover {
                background-color: #f3f4f6;
                color: #000000;
            }
        """)
        box.setStyleSheet(f"""
            QMessageBox {{
                background: {self.theme['surface']};
            }}
            QMessageBox QLabel {{
                color: {self.theme['text']};
                font-size: 13px;
            }}
        """)
        box.exec_()

    def show_info(self, title, text):
        box = QMessageBox(self)
        box.setIcon(QMessageBox.Information)
        box.setWindowTitle(title)
        box.setText(text)
        box.setStandardButtons(QMessageBox.Ok)
        box.setStyleSheet(f"""
            QMessageBox {{
                background: {self.theme['surface']};
            }}
            QMessageBox QLabel {{
                color: {self.theme['text']};
                font-size: 13px;
            }}
            QPushButton {{
                min-width: 90px;
                min-height: 34px;
                border-radius: 10px;
                border: 1px solid #cfe0ff;
                background: #edf3ff;
                color: {self.theme['accent']};
                font-weight: 600;
                padding: 0 12px;
            }}
            QPushButton:hover {{
                background: #e4edff;
            }}
        """)
        box.exec_()

    def _on_summary_finished(self, text):
        self.sum_output.setText(text)
        if text.lower().startswith("summarization error"):
            self.set_status_message("Summary failed. Please try again.", "error")
            self.header_status.setText("Error")
        else:
            self.set_status_message("Summary completed", "success")
            self.header_status.setText("Done")

    def _on_generation_finished(self, text):
        self.gen_output.setText(text)
        self.last_generated = text
        if text.lower().startswith("generation error"):
            self.set_status_message("Generation failed. Please try again.", "error")
            self.header_status.setText("Error")
        else:
            self.set_status_message("Document generated", "success")
            self.header_status.setText("Done")

    # --------------------------------------------------
    # LOGIC
    # --------------------------------------------------
    def select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Folder")
        if folder:
            folder_name = folder.split("/")[-1] or folder.split("\\")[-1]
            self.folder_name.setText(folder_name)
            self.footer_folder.setText(folder_name)
            self.set_status_message("Scanning documents...", "muted")
            self.header_status.setText("Scanning...")

            self.documents = scan_folder(folder)

            self.set_status_message(f"{len(self.documents)} document(s) loaded", "success")
            self.header_status.setText(f"{len(self.documents)} loaded")

            self.load_deadlines()
            self.refresh_summarize_documents()

    def load_deadlines(self):
        self.deadline_list.clear()
        found = False

        for doc in self.documents:
            if doc["extension"] == ".xlsx":
                alerts = check_expiry_from_excel(doc["filepath"], days_threshold=30)
                for a in alerts:
                    found = True
                    item = QListWidgetItem(
                        f"{a['employee']}  •  {a['date']}  •  {a['days_left']} day(s) left"
                    )
                    item.setForeground(QColor(self.theme["danger"]))
                    # Attach filepath so dragging/copying the source spreadsheet works.
                    item.setData(Qt.UserRole, doc)
                    self.deadline_list.addItem(item)

        if not found:
            item = QListWidgetItem("No deadlines within the next 30 days")
            item.setForeground(QColor(self.theme["success"]))
            self.deadline_list.addItem(item)

    def search(self):
        query = self.search_input.text().strip()

        if not query:
            self.show_warning("Please enter a search query.")
            return

        if not self.documents:
            self.show_warning("Please select a folder first.")
            return

        results = search_documents(self.documents, query)
        self.search_results = results
        self.results_list.clear()
        self.snippet_box.clear()

        self.set_results_count(len(results))

        if results:
            for r in results:
                ext = r['filename'].split('.')[-1].lower()
                if ext == 'pdf':
                    icon = "📕"
                elif ext == 'docx':
                    icon = "📘"
                elif ext == 'xlsx':
                    icon = "📗"
                else:
                    icon = "📄"
                match_icon = "🔍" if r.get('match_type') == 'semantic' else ""
                item = QListWidgetItem(f"{icon}  {r['filename']}  {match_icon}")
                item.setData(Qt.UserRole, r)
                self.results_list.addItem(item)

        self.update_empty_states()

    def show_snippet(self, item):
        index = self.results_list.row(item)

        if index < len(self.search_results):
            text = self.search_results[index]["snippet"]
            query = self.search_input.text().strip()
            t = self.theme

            html = text.replace("\n", "<br>")

            if query:
                html = re.sub(
                    f"({re.escape(query)})",
                    (
                        f"<span style='background:{t['accent_soft']}; color:{t['text']}; "
                        f"padding:2px 5px; border-radius:4px;'>\\1</span>"
                    ),
                    html,
                    flags=re.IGNORECASE
                )

            self.snippet_box.setHtml(
                f"""
                <div style="
                    font-family:'Segoe UI';
                    font-size:13px;
                    color:{t['text']};
                    line-height:1.7;
                ">
                    {html}
                </div>
                """
            )

        self.update_empty_states()

    def summarize_selected(self):
        item = self.sum_doc_list.currentItem()
        if item is None:
            self.show_warning("Please select a document.")
            return

        doc = item.data(Qt.UserRole)
        if not doc:
            self.show_warning("Selected document is invalid.")
            return

        self.sum_output.setText("Summarizing... Please wait.")
        self.set_status_message("AI is preparing summary...", "muted")
        self.header_status.setText("Working...")

        self.worker = WorkerThread(lambda d: summarize_document(d)["summary"], doc)
        self.worker.finished.connect(self._on_summary_finished)
        self.worker.start()

    def generate_doc(self):
        details = self.gen_input.text().strip()

        if not details:
            self.show_warning("Please enter document details.")
            return

        doc_type = self.doc_type.currentText()
        self.gen_output.setText("Generating... Please wait.")
        self.set_status_message("AI is generating document...", "muted")
        self.header_status.setText("Working...")

        self.worker = WorkerThread(generate_document, doc_type, details)
        self.worker.finished.connect(self._on_generation_finished)
        self.worker.start()

    def save_as_word(self):
        if not self.last_generated:
            self.show_warning("No generated content to save.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self,
            "Save File",
            "",
            "Word Files (*.docx)"
        )

        if path:
            doc = DocxDocument()
            doc.add_heading("Generated Document", 0)

            for line in self.last_generated.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)

            doc.save(path)
            self.show_info("Saved", f"Saved:\n{path}")
            self.set_status_message(f"Saved: {path}", "success")
            self.header_status.setText("Saved")
            self.header_status.setText("Saved")


class FileCopyDragListWidget(QListWidget):
    """
    Drag items out to Desktop/Explorer as COPY (no delete/move).
    We expose local file paths via mime URLs so dropping creates copies.
    """

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.setDragEnabled(True)
        self.setDragDropMode(QAbstractItemView.DragOnly)
        self._drag_start_pos = None

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._drag_start_pos = event.pos()
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if not (event.buttons() & Qt.LeftButton):
            self._drag_start_pos = None
            return super().mouseMoveEvent(event)

        if self._drag_start_pos is None:
            return super().mouseMoveEvent(event)

        if (event.pos() - self._drag_start_pos).manhattanLength() < QApplication.startDragDistance():
            return super().mouseMoveEvent(event)

        selected_items = self.selectedItems()
        item_under = self.itemAt(event.pos())
        if item_under is not None and item_under not in selected_items:
            selected_items = [item_under]

        file_paths = []
        for it in selected_items:
            payload = it.data(Qt.UserRole)
            if isinstance(payload, dict):
                p = payload.get("filepath") or payload.get("path")
            else:
                p = str(payload) if payload is not None else None

            if p and os.path.exists(p):
                file_paths.append(p)

        if not file_paths:
            self._drag_start_pos = None
            return

        # De-duplicate while keeping order.
        seen = set()
        unique_paths = []
        for p in file_paths:
            if p not in seen:
                unique_paths.append(p)
                seen.add(p)

        mime = QMimeData()
        mime.setUrls([QUrl.fromLocalFile(p) for p in unique_paths])

        drag = QDrag(self)
        drag.setMimeData(mime)
        drag.exec_(Qt.CopyAction)

        self._drag_start_pos = None

