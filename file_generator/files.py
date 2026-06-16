import os
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt

# Different data sets to generate unique test files
SCENARIOS = [
    {
        "id": "01",
        "company_a": 'LLC "DressUp Georgia"',
        "company_b": 'LLC "Global Logistics"',
        "contract_no": "2026/04-01",
        "period": "Q1 2026",
        "city": "Tbilisi",
        "sales_rows": [
            ["2026-01-10", "Adidas", "Shoes", "Ultraboost Light", "Recycled Polyester", 12, 430],
            ["2026-01-16", "Nike", "Apparel", "Tech Fleece", "Cotton", 20, 315],
            ["2026-02-03", "Puma", "Accessories", "Gym Bag", "Synthetic", 35, 82],
            ["2026-03-01", "Reebok", "Shoes", "Nano X3", "Textile", 9, 390],
        ],
    },
    {
        "id": "02",
        "company_a": 'LLC "Nova Sport"',
        "company_b": 'LLC "Euro Freight"',
        "contract_no": "2026/04-02",
        "period": "Q2 2026",
        "city": "Batumi",
        "sales_rows": [
            ["2026-04-04", "Under Armour", "Apparel", "HeatGear Tee", "Polyester", 40, 72],
            ["2026-04-18", "Nike", "Shoes", "Air Zoom Pegasus", "Mesh", 16, 355],
            ["2026-05-07", "Adidas", "Accessories", "Cap", "Cotton", 55, 35],
            ["2026-06-14", "Puma", "Shoes", "Velocity Nitro", "Synthetic", 10, 410],
        ],
    },
    {
        "id": "03",
        "company_a": 'LLC "FitGear"',
        "company_b": 'LLC "Caucasus Logistics"',
        "contract_no": "2026/04-03",
        "period": "Q3 2026",
        "city": "Kutaisi",
        "sales_rows": [
            ["2026-07-03", "Hummel", "Apparel", "Training Shorts", "Polyester", 28, 68],
            ["2026-07-19", "Asics", "Shoes", "Gel-Kayano", "Textile", 11, 520],
            ["2026-08-25", "Mizuno", "Shoes", "Wave Rider", "Textile", 14, 470],
            ["2026-09-09", "Nike", "Accessories", "Backpack", "Nylon", 33, 95],
        ],
    },
    {
        "id": "04",
        "company_a": 'LLC "Sportline"',
        "company_b": 'LLC "Transport Hub"',
        "contract_no": "2026/04-04",
        "period": "Q4 2026",
        "city": "Rustavi",
        "sales_rows": [
            ["2026-10-05", "Adidas", "Shoes", "Gazelle", "Leather", 18, 280],
            ["2026-10-21", "Puma", "Apparel", "Runner Jacket", "Polyester", 22, 140],
            ["2026-11-11", "New Balance", "Shoes", "574", "Suede", 17, 360],
            ["2026-12-02", "Nike", "Apparel", "Dri-FIT Set", "Cotton/Polyester", 25, 210],
        ],
    },
]


def ensure_output_dir():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(base_dir, "generated_test_files")
    os.makedirs(output_dir, exist_ok=True)
    return output_dir


def create_contract_pdf(scenario, output_dir):
    pdf_path = os.path.join(output_dir, f"contract_{scenario['id']}.pdf")
    lines = [
        f"Partnership Agreement #{scenario['contract_no']}",
        "",
        f"City: {scenario['city']}",
        f"Party A: {scenario['company_a']}",
        f"Party B: {scenario['company_b']}",
        "",
        "1) Scope: logistics and transportation services.",
        "2) Payment: within 5 banking days from invoice.",
        "3) Liability: full compensation for delay or damage.",
        "",
        "Signature (Party A): ______________________",
        "Signature (Party B): ______________________",
    ]
    write_pdf_unicode(pdf_path, lines)
    return pdf_path


def create_excel(scenario, output_dir):
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Sales Report"

    headers = [
        "Date",
        "Brand",
        "Category",
        "Product Name",
        "Material",
        "Quantity",
        "Unit Price",
        "Total Value",
    ]
    ws1.append(headers)
    for idx, row in enumerate(scenario["sales_rows"], start=2):
        ws1.append(row + [f"=F{idx}*G{idx}"])

    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws1[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    for col in range(1, len(headers) + 1):
        ws1.column_dimensions[get_column_letter(col)].width = 20

    ws2 = wb.create_sheet("Financial Summary")
    ws2.append(["Brand", "Total Sales"])
    brands = sorted({row[1] for row in scenario["sales_rows"]})
    for i, brand in enumerate(brands, start=2):
        ws2.cell(row=i, column=1, value=brand)
        ws2.cell(row=i, column=2, value=f"=SUMIF('Sales Report'!B:B, A{i}, 'Sales Report'!H:H)")

    ws2.cell(row=len(brands) + 3, column=1, value="Grand Total")
    ws2.cell(row=len(brands) + 3, column=2, value=f"=SUM(B2:B{len(brands) + 1})")

    excel_path = os.path.join(output_dir, f"fin_report_{scenario['id']}.xlsx")
    wb.save(excel_path)
    return excel_path


def create_financial_pdf(scenario, output_dir):
    total_revenue = sum(row[5] * row[6] for row in scenario["sales_rows"])
    operating_cost = int(total_revenue * 0.35)
    net_profit = total_revenue - operating_cost
    pdf_path = os.path.join(output_dir, f"financial_report_{scenario['id']}.pdf")
    lines = [
        f"Quarterly Financial Report ({scenario['period']})",
        "",
        f"Total Revenue: {total_revenue:,} GEL",
        f"Operating Cost: {operating_cost:,} GEL",
        f"Net Profit: {net_profit:,} GEL",
        "",
        "Brand Performance:",
    ]
    for row in scenario["sales_rows"]:
        lines.append(f"- {row[1]} | Quantity: {row[5]} | Unit Price: {row[6]} GEL")
    lines.append("")
    lines.append("Document prepared for QA testing.")
    write_pdf_unicode(pdf_path, lines)
    return pdf_path


def get_unicode_font_path():
    candidate_paths = [
        "C:\\Windows\\Fonts\\arial.ttf",
        "C:\\Windows\\Fonts\\segoeui.ttf",
        "C:\\Windows\\Fonts\\calibri.ttf",
        "C:\\Windows\\Fonts\\times.ttf",
        "C:\\Windows\\Fonts\\DejaVuSans.ttf",
    ]
    for path in candidate_paths:
        if os.path.exists(path):
            return path
    raise FileNotFoundError("Unicode font not found in C:\\Windows\\Fonts")


def write_pdf_unicode(path, lines):
    font_path = get_unicode_font_path()
    font_name = "UnicodeFont"
    if font_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(font_name, font_path))

    pdf = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    y = height - 50
    pdf.setFont(font_name, 12)

    for line in lines:
        pdf.drawString(50, y, line)
        y -= 18
        if y < 50:
            pdf.showPage()
            pdf.setFont(font_name, 12)
            y = height - 50
    pdf.save()


def create_agreement_details_pdf(scenario, output_dir):
    pdf_path = os.path.join(output_dir, f"agreement_details_{scenario['id']}.pdf")
    lines = [
        f"Agreement Details #{scenario['contract_no']}",
        "",
        f"Company 1: {scenario['company_a']}",
        f"Company 2: {scenario['company_b']}",
        "Service Type: delivery and warehouse logistics",
        "Delivery Window: 7-10 business days",
        "Invoice Deadline: within 24 hours after delivery",
        "Penalty on Breach: 2% of contract value",
        "",
        "Document generated for test data purposes.",
    ]
    write_pdf_unicode(pdf_path, lines)
    return pdf_path


def create_word_agreement(scenario, output_dir):
    doc = Document()
    title = doc.add_heading(f"Partnership Agreement #{scenario['contract_no']}", level=1)
    title.alignment = 1

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_paragraph(f"This document is prepared in {scenario['city']}.")
    doc.add_paragraph(f"Parties: {scenario['company_a']} and {scenario['company_b']}.")

    doc.add_heading("Key Terms", level=2)
    doc.add_paragraph("1) Delivery window: 7-10 business days.")
    doc.add_paragraph("2) Payment terms: within 5 banking days from invoice.")
    doc.add_paragraph("3) Liability: full compensation for delay or damage.")
    doc.add_paragraph("4) Disputes resolved under applicable law.")

    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Service"
    hdr[1].text = "Monthly Volume"
    hdr[2].text = "Price (GEL)"

    for service, volume, price in [("Transportation", "12 deliveries", "24,000"), ("Insurance", "12 policies", "6,000")]:
        row = table.add_row().cells
        row[0].text = service
        row[1].text = volume
        row[2].text = price

    doc.add_paragraph("\nSignature A: __________________    Signature B: __________________")
    word_path = os.path.join(output_dir, f"agreement_{scenario['id']}.docx")
    doc.save(word_path)
    return word_path


def generate_all_test_files():
    output_dir = ensure_output_dir()
    generated_files = []

    for scenario in SCENARIOS:
        generated_files.append(create_excel(scenario, output_dir))
        generated_files.append(create_contract_pdf(scenario, output_dir))
        generated_files.append(create_financial_pdf(scenario, output_dir))
        generated_files.append(create_word_agreement(scenario, output_dir))

    return generated_files, output_dir

if __name__ == "__main__":
    files, output_folder = generate_all_test_files()
    print(f"Generated {len(files)} files in: {output_folder}")
    for file_path in files:
        print(f"- {os.path.basename(file_path)}")